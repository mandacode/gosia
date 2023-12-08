import typing as tp

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.conf import settings


def add_contacts_table(
        document: Document,
        recipient_info: str,
        sender_info: str,
        recipient_header: tp.Optional[str] = None,
        sender_header: tp.Optional[str] = None
):
    rows = 2
    cols = 2
    table = document.add_table(rows=rows, cols=cols)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

    if recipient_header:
        table.cell(0, 0).text = recipient_header
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    if sender_header:
        table.cell(0, 1).text = sender_header
        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph = table.cell(1, 0).add_paragraph()
    run = paragraph.add_run(recipient_info)
    run.bold = True
    run.font.size = Pt(14)
    table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = table.cell(1, 1).add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(sender_info)
    run.bold = True
    run.font.size = Pt(14)
    table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_header(document: Document, invoice_header: str, invoice_location_and_date: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(invoice_header)
    run.bold = True
    run.font.size = Pt(12)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(invoice_location_and_date)


def add_text(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_worked_hours_table_for_employee(document: Document, worked_hours: list):
    rows = len(worked_hours) + 1
    cols = 4
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    columns = ['POS.', 'DATUM', 'MENGE', 'BETRAG']
    for idx, name in enumerate(columns):
        paragraph = table.cell(0, idx).paragraphs[0]
        run = paragraph.add_run(name)
        run.bold = True

    for row_idx, row in enumerate(worked_hours):
        for col_idx, col in enumerate(row):
            cell = table.cell(row_idx + 1, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(col)


def add_worked_hours_table_for_customer(document: Document, worked_hours: list):
    rows = len(worked_hours) + 1
    cols = 3
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    columns = ['Datum', 'Stunden', 'Betrag']
    for idx, name in enumerate(columns):
        paragraph = table.cell(0, idx).paragraphs[0]
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif idx == 1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(name)
        run.bold = True

    for row_idx, row in enumerate(worked_hours):
        for col_idx, col in enumerate(row):
            cell = table.cell(row_idx + 1, col_idx)
            paragraph = cell.paragraphs[0]
            if col_idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif col_idx == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if col_idx == 0:
                run = paragraph.add_run(f'{row_idx + 1}. ')
                run.bold = True
            paragraph.add_run(col)

    document.add_paragraph('_' * 82)


def add_employee_total(document: Document, total_text: str, total: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(total_text)
    run = paragraph.add_run(total)
    run.bold = True
    run.font.size = Pt(12)


def add_customer_total(document: Document, netto: str, tax: str, brutto: str):
    rows = 3
    cols = 5
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell = table.cell(0, 3)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run('Nettobetrag')

    cell = table.cell(0, 4)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(netto)

    cell = table.cell(1, 3)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run('zzgl.19%  MwSt.')

    cell = table.cell(1, 4)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(tax)

    cell = table.cell(2, 3)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run('Gesamtbetrag')

    cell = table.cell(2, 4)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(brutto)
    run.bold = True
    run.font.size = Pt(14)


def add_signature(document: Document, signature: str, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    paragraph.add_run(signature)


def add_footer(document: Document, footer: str, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    paragraph.add_run(footer)


def generate_employee_invoice(data: dict):
    document = Document()
    document.styles['Normal'].font.name = 'Arial'
    document.styles['Normal'].font.size = Pt(10)

    add_contacts_table(
        document=document,
        recipient_header=data['invoice_recipient_header'],
        sender_header=data['invoice_sender_header'],
        recipient_info=data['invoice_recipient'],
        sender_info=data['invoice_sender']
    )
    add_header(
        document=document,
        invoice_header=data['invoice_header'],
        invoice_location_and_date=data['invoice_location_and_date']
    )
    add_text(document=document, text=data['invoice_text'])
    add_worked_hours_table_for_employee(
        document=document,
        worked_hours=data['invoice_worked_hours'],
    )
    add_employee_total(
        document=document,
        total_text=data['invoice_total_text'],
        total=data['invoice_total']
    )
    add_signature(document=document, signature=data['invoice_signature'])
    add_footer(document=document, footer=data['invoice_footer'])
    document.save("employee_invoice_example.docx")


def generate_customer_invoice(data: dict):
    document = Document(settings.BASE_DIR / "invoices" / "generator" / "templates" / "customer_invoice.docx")
    add_contacts_table(
        document=document,
        recipient_info=data['recipient'],
        sender_info=data['sender']
    )
    add_header(
        document=document,
        invoice_header=data['header'],
        invoice_location_and_date=data['location_and_date']
    )
    add_text(document=document, text=data['text'])
    add_worked_hours_table_for_customer(
        document=document,
        worked_hours=data['worked_hours'],
    )
    add_customer_total(
        document=document,
        netto=data['netto'],
        tax=data['tax'],
        brutto=data['brutto'],
    )
    add_signature(document=document, signature=data['signature'], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_footer(document=document, footer=data['footer'])
    document.save(settings.BASE_DIR / "static" / "docs" / "customer_invoice_example.docx")


def main():

    employee_data = {
        "invoice_number": "2023/05",
        "invoice_date": "26.11.2023",
        "invoice_header": "Rechnung Nr: 2023/05",
        "invoice_recipient_header": "AUFTRAGGEBER",
        "invoice_sender_header": "AUFTRAGNEHMER",
        "invoice_location_and_date": "Berlin, den 31.10.2023",
        "invoice_recipient": "Reinigungsservice\nJohn Doe\nFakestreet\n1123 FakeCity\nDE 123",
        "invoice_sender": "\nAnna Nowak\nFakestreet\nFakeCity/Fakecountry\nSt.Nr. 123123123",
        "invoice_text": "Sehr geehrte Frau Jarmul,\n\nIch badanke mich für Ihren Auftrag und für meine Reinigungsarbeiten stelle ich Ihnen folgenden Positionen in Rechnung. Der vereinbarte Stundenlohn beträgt 12,-€.",
        "invoice_worked_hours": [
            ['1.', '01.11.2023', '10 Std', '120,00 €'],
            ['2.', '02.11.2023', '9 Std', '108,00 €'],
            ['3.', '03.11.2023', '12 Std', '144,00 €'],
            ['4.', '03.11.2023', '12 Std', '144,00 €'],
            ['5.', '03.11.2023', '12 Std', '144,00 €'],
            ['6.', '03.11.2023', '12 Std', '144,00 €'],
            ['7.', '03.11.2023', '12 Std', '144,00 €'],
            ['8.', '03.11.2023', '12 Std', '144,00 €'],
            ['9.', '03.11.2023', '12 Std', '144,00 €'],
            ['10.', '03.11.2023', '12 Std', '144,00 €']
        ],
        "invoice_total_text": "Ich bitte um Überweisung des Gesamtbetrages von:" + 30 * " ",
        "invoice_total": "2472,00 €",
        "invoice_signature": "Gesamtbetrag dankend in bar erhalten",
        "invoice_footer": "Mit freundlichen Grüßen\n\nIvanna Chornysh\n\nHinweis;  Wegen Kleingewerberegelung ist der Auftragnehmer von der Umsatzsteuer befreit."
    }

    customer_data = {
        "invoice_number": "2023/919",
        "invoice_date": "26.11.2023",
        "invoice_header": "Rechnungsnummer: 2023/919",
        "invoice_location_and_date": "Berlin, den 26.11.2023",
        "invoice_recipient": "\nAnna Nowak\nFakestreet\nFakeCity/Fakecountry",
        "invoice_sender": "John Doe\nFakestreet\n1123 FakeCity\nDE 123",
        "invoice_text": "Sehr geehrte Frau Eilers,\n\nwir haben für Sie Reinigungsarbeiten in Ihrem Haushalt erbracht. Die Abrechnung entnehmen Sie bitte der beigefügten Aufstellung",
        "invoice_worked_hours": [
            ['03.10.2023', '6 Std', '78,00 €'],
            ['10.10.2023', '5 Std', '65,00 €'],
            ['12.10.2023', '6 Std', '78,00 €'],
            ['17.10.2023', '5 Std', '65,00 €'],
            ['24.10.2023', '3 Std', '39,00 €'],
            ['27.10.2023', '4 Std', '52,00 €'],
            ['31.10.2023', '4 Std', '52,00 €'],
        ],
        "invoice_total_netto": "429,00 €",
        "invoice_tax": "81,51 €",
        "invoice_total_brutto": "510,51 €",
        "invoice_signature": "Bitte überweisen Sie den Gesamtbetrag innerhalb von 10 Tagen auf folgendes Konto:\n\nBank: Berliner Sparkasse\nIBAN: DE 12 1001 0010 0987 6543 21\nBIC: PBNKDEFFXXX\n\nVielen Dank für Ihren Auftrag!",
        "invoice_footer": "Mit freundlichen Grüßen\n\nMalgorzata Jarmul"
    }

    generate_employee_invoice(data=employee_data)
    generate_customer_invoice(data=customer_data)


if __name__ == '__main__':
    main()
