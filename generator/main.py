from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


FILENAME = 'test.docx'

def add_contacts_table(
        document: Document, recipient_header: str, sender_header: str, recipient_info: str, sender_info: str
):
    rows = 2
    cols = 3
    table = document.add_table(rows=rows, cols=cols)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Set to white

    table.cell(0, 0).text = recipient_header
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    table.cell(0, 2).text = sender_header
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    paragraph = table.cell(1, 0).add_paragraph()
    run = paragraph.add_run(recipient_info)
    run.bold = True
    table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph = table.cell(1, 2).add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(sender_info)
    run.bold = True
    table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_header(document: Document, invoice_header: str, invoice_location_and_date: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(invoice_header)
    run.bold = True

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(invoice_location_and_date)


def add_text(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_worked_hours_table(document: Document, worked_hours: list):
    rows = len(worked_hours) + 1
    cols = 4
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    columns = ['POS.', 'DATUM', 'MENGE', 'BETRAG']
    for idx, name in enumerate(columns):
        paragraph = table.cell(0, idx).paragraphs[0]
        run = paragraph.add_run(name)
        run.bold = True

    for row_idx, row in enumerate(worked_hours):
        for col_idx, col in enumerate(row):
            cell = table.cell(row_idx + 1, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(col)


def add_total(document: Document, total_text: str, total: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(total_text)
    run = paragraph.add_run(total)
    run.bold = True


def add_signature(document: Document, signature: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run(signature)


def add_footer(document: Document, footer: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(footer)


def main():

    data = {
        "invoice_number": "2023/05",
        "invoice_date": "26.11.2023",
        "invoice_header": "Rechnung Nr: 2023/05",
        "invoice_recipient_header": "AUFTRAGGEBER",
        "invoice_sender_header": "AUFTRAGNEHMER",
        "invoice_location_and_date": "Berlin, den 26.11.2023",
        "invoice_recipient": "Reinigungsservice\nMalgorzata Jarmul\nHauptstr. 100\n10 827 Berlin\nDE 311547701",
        "invoice_sender": "\nIvanna Chornysh\nPeremogi 50\n81000 Javoriv/Ukraine\nSt.Nr. 3289202826",
        "invoice_text": "Sehr geehrte Frau Jarmul,\n\nIch badanke mich für Ihren Auftrag und für meine Reinigungsarbeiten stelle ich Ihnen folgenden Positionen in Rechnung. Der vereinbarte Stundenlohn beträgt 12,-€.",
        "invoice_worked_hours": [
            ['1.', '01.11.2023', '10 Std', '120,00 €'],
            ['2.', '02.11.2023', '9 Std', '108,00 €'],
            ['3.', '03.11.2023', '12 Std', '144,00 €'],
            ['4.', '03.11.2023', '12 Std', '144,00 €'],
            ['5.', '03.11.2023', '12 Std', '144,00 €'],
            ['6.', '03.11.2023', '12 Std', '144,00 €'],
            ['7.', '03.11.2023', '12 Std', '144,00 €'],
            ['8.', '03.11.2023', '12 Std', '144,00 €'],
            ['9.', '03.11.2023', '12 Std', '144,00 €'],
            ['10.', '03.11.2023', '12 Std', '144,00 €']
        ],
        "invoice_total_text": "Ich bitte um Überweisung des Gesamtbetrages von:" + 30 * " ",
        "invoice_total": "2472,00 €",
        "invoice_signature": "Gesamtbetrag dankend in bar erhalten",
        "invoice_footer": "Mit freundlichen Grüßen\n\nIvanna Chornysh\n\nHinweis;  Wegen Kleingewerberegelung ist der Auftragnehmer von der Umsatzsteuer befreit."
    }

    # Create a new Word document
    document = Document()
    add_contacts_table(
        document=document,
        recipient_header=data['invoice_recipient_header'],
        sender_header=data['invoice_sender_header'],
        recipient_info=data['invoice_recipient'],
        sender_info=data['invoice_sender']
    )
    add_header(
        document=document,
        invoice_header=data['invoice_header'],
        invoice_location_and_date=data['invoice_location_and_date']
    )
    add_text(document=document, text=data['invoice_text'])
    add_worked_hours_table(
        document=document,
        worked_hours=data['invoice_worked_hours'],
    )
    add_total(
        document=document,
        total_text=data['invoice_total_text'],
        total=data['invoice_total']
    )
    add_signature(document=document, signature=data['invoice_signature'])
    add_footer(document=document, footer=data['invoice_footer'])
    document.save(FILENAME)


if __name__ == '__main__':
    main()
